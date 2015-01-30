# -*- coding: utf-8 -*-
'''
Created on 2015/1/2

@author: sam
'''

import re
import xlrd
import docx
from itertools import izip_longest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT 
from wx import ITEM_CHECK

def itoa(num):
    return str(int(num)) if isinstance(num,(int,long,float)) else num

def parseSheet(sheet):
    parseResult={}
    pDate = re.compile("^(\d{4,4})/(\d{1,2})/(\d{1,2})")
    pId = re.compile("^\d{7,7}$")
    for i in range(0,sheet.nrows):
        m1 = re.match(pDate,sheet.row_values(i)[0])
        if m1 :
            date = "{0:d}/{1:02d}/{2:02d}".format(*map(int,m1.groups()))
            ptlist = parseResult.setdefault(date,[])
        item4 = sheet.row_values(i)[4]
        item5 = sheet.row_values(i)[5]
        id4 = itoa(item4)
        id5 = itoa(item5)
        m4 = re.match(pId,id4)
        m5 = re.match(pId,id5)
        if m4 :
            info = sheet.row_values(i,1)
            ptlist.append(info)
        elif m5:
            info = sheet.row_values(i,1)
            info = info[0:1]+ [info[1]+"\n"+info[2]] + info[3:]
            ptlist.append(info)
    return parseResult

def genCardString(pt):
    if pt:
        s = ( pt[0] + "\n\n"
            + "\n".join(map(itoa ,pt[3:6]))
            + "," + pt[6] + "\n\n"
            + "\n\n".join(map(itoa ,pt[7:]))
            + pt[2] )
    else:
        s = ""
    return s

def writeCell(cell,pt):
    paragraph = cell.paragraphs[0]
    paragraph.style = "BodyText3"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = genCardString(pt)

def writeRow(row,row_info):
    # row_info: list of pt_info of length 6
    pt_list = row_info
    cells = row.cells
    map(lambda cell,pt : writeCell(cell,pt), cells, pt_list)

def addTable(doc,table_info):
    #table_info: list of pt_info of length 3x6
    table = doc.add_table(3,6,style="a3")
    table.autofit=False
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    row_info_iter = izip_longest(*[iter(table_info)]*6,fillvalue=[])
    map(lambda row, row_info: writeRow(row, row_info), table.rows, row_info_iter)
    for column in table.columns:
        column.width = docx.shared.Cm(3)

    

def makeCard(doc,ptlist):
    table_info_iter = izip_longest(*[iter(ptlist)]*18,fillvalue=[])
    map( lambda table_info: addTable(doc,table_info) , table_info_iter)

def docsetup(doc):
    section = doc.sections[0]
    section.top_margin = docx.shared.Cm(1)
    section.bottom_margin = docx.shared.Cm(1)
    section.left_margin = docx.shared.Cm(1)
    section.right_margin = docx.shared.Cm(1)
    section.page_width = docx.shared.Cm(21)
    section.page_height = docx.shared.Cm(29.7)

def getptlist(filename,sheetname,datelabel):
    pDate = re.compile("^(\d{4,4})/(\d{1,2})/(\d{1,2})")
    m = re.match(pDate,datelabel)
    date = "{0:4}/{1:02}/{2:02}".format(*map(int,m.groups()))
    book = xlrd.open_workbook(filename)
    sheet = book.sheet_by_name(sheetname)
    parseResult = parseSheet(sheet)
    ptlist = parseResult.get(date,[])
    ptlist = [[datelabel]+x for x in ptlist]
    return ptlist

def main(filenamelist,sheetname,outputfile,datelabel):
    doc = docx.Document("templates/Doc1.docx")
    paragraph = doc.add_paragraph(text=datelabel+u" 入院列表",style="a4")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docsetup(doc)
    ptlist = sum(map(lambda filename: getptlist(filename,sheetname,datelabel),filenamelist),[])
    makeCard(doc,ptlist)
    doc.save(outputfile)
    return u"產生清單成功！！"

if __name__ == '__main__':
    main(["../Cath20150102.xls","../Hema20150102.xls","../others20150102.xls"],"201501","2015/1/5(W1)")
    